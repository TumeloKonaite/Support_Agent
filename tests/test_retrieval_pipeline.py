import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from src.app.core.config import Settings
from src.app.core.dependencies import get_retriever
from src.app.domain.support.retrieval import RetrievalPipeline
from src.app.infrastructure.retrieval.indexer import (
    KnowledgeDocument,
    build_indexer,
    run_query,
)
from src.app.infrastructure.retrieval.retriever import VectorStoreRetriever
from src.app.infrastructure.retrieval.vector_store import JsonVectorStore


class RetrievalPipelineTests(unittest.TestCase):
    def test_chunk_documents_applies_overlap(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            vector_store_path = data_dir / "vector_store.json"
            indexer = build_indexer(
                Settings(
                    content_data_dir=data_dir,
                    retrieval_vector_store_path=vector_store_path,
                    retrieval_chunk_size=30,
                    retrieval_chunk_overlap=10,
                )
            )

            chunks = indexer.chunk_documents(
                [
                    KnowledgeDocument(
                        source="memory",
                        content="abcdefghijklmnopqrstuvwxyz1234567890",
                        metadata={"kind": "test"},
                    )
                ]
            )

        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0].text, "abcdefghijklmnopqrstuvwxyz1234")
        self.assertEqual(chunks[1].text, "uvwxyz1234567890")

    def test_indexer_loads_json_knowledge_and_persists_vectors(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            self._write_json(
                data_dir / "knowledge.json",
                {
                    "sections": {
                        "Policies": [
                            "Escalate refund exceptions to a human specialist.",
                        ]
                    }
                },
            )
            vector_store_path = data_dir / "retrieval" / "vectors.json"
            indexer = build_indexer(
                Settings(
                    content_data_dir=data_dir,
                    retrieval_vector_store_path=vector_store_path,
                    retrieval_chunk_size=80,
                    retrieval_chunk_overlap=20,
                )
            )

            chunks = indexer.index()
            store = JsonVectorStore(vector_store_path)
            persisted_count = store.count()
            exists = vector_store_path.exists()

        self.assertEqual(len(chunks), 1)
        self.assertEqual(persisted_count, 1)
        self.assertTrue(exists)

    def test_indexer_loads_docx_blocks_and_tables_with_metadata(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            docx_path = data_dir / "policies" / "refund-policy.docx"
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            document = Document()
            document.add_heading("Cancellation and Refund Policy", level=1)
            document.add_paragraph("Refund requests are reviewed by the support team.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Window"
            table.cell(0, 1).text = "Action"
            table.cell(1, 0).text = "14 days"
            table.cell(1, 1).text = "Escalate to human support"
            document.save(docx_path)
            vector_store_path = data_dir / "retrieval" / "vectors.json"
            indexer = build_indexer(
                Settings(
                    content_data_dir=data_dir,
                    retrieval_vector_store_path=vector_store_path,
                    retrieval_chunk_size=160,
                    retrieval_chunk_overlap=20,
                )
            )

            chunks = indexer.index()
            store = JsonVectorStore(vector_store_path)
            persisted_count = store.count()

        self.assertGreaterEqual(len(chunks), 4)
        self.assertEqual(persisted_count, len(chunks))
        self.assertTrue(any("Cancellation and Refund Policy" in chunk.text for chunk in chunks))
        table_chunks = [chunk for chunk in chunks if "14 days | Escalate" in chunk.text]
        self.assertEqual(len(table_chunks), 1)
        self.assertEqual(table_chunks[0].metadata["file_type"], "docx")
        self.assertEqual(table_chunks[0].metadata["filename"], "refund-policy.docx")
        self.assertEqual(table_chunks[0].metadata["docx_block_type"], "TableRow")

    def test_retrieval_returns_relevant_business_knowledge(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            self._write_json(
                data_dir / "business_profile.json",
                {
                    "business_name": "Glow Studio",
                    "support_hours": "Monday to Friday, 9:00 AM to 5:00 PM local time.",
                },
            )
            self._write_json(
                data_dir / "knowledge.json",
                {
                    "sections": {
                        "Policies": [
                            "Escalate refund requests to the human support team.",
                            "Ask a clarifying question if the request is ambiguous.",
                        ],
                        "FAQs": [
                            "Customers often ask about products and support hours.",
                        ],
                    }
                },
            )
            settings = Settings(
                content_data_dir=data_dir,
                retrieval_vector_store_path=data_dir / "retrieval" / "vectors.json",
                retrieval_top_k=2,
                retrieval_chunk_size=120,
                retrieval_chunk_overlap=20,
            )
            indexer = build_indexer(settings)
            indexer.index()

            results = run_query(indexer, "When are your support hours?", top_k=2)

        self.assertEqual(len(results), 2)
        self.assertIn("support hours", results[0].chunk.text.lower())

    def test_retriever_overfetches_before_lexical_reranking(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            self._write_json(
                data_dir / "business_profile.json",
                {
                    "business_name": "Glow Studio",
                    "support_hours": "Monday to Friday, 9:00 AM to 5:00 PM local time.",
                    "support_phone": "+1-800-555-0130",
                },
            )
            self._write_json(
                data_dir / "knowledge.json",
                {
                    "sections": {
                        "FAQs": [
                            "WHAT ARE COOKIES?",
                            "Burns;",
                            "Support phone questions can be routed to the care team.",
                        ],
                    },
                },
            )
            settings = Settings(
                content_data_dir=data_dir,
                retrieval_vector_store_path=data_dir / "retrieval" / "vectors.json",
                retrieval_top_k=1,
                retrieval_chunk_size=120,
                retrieval_chunk_overlap=20,
            )
            indexer = build_indexer(settings)
            indexer.index()

            results = run_query(indexer, "What are your support hours?", top_k=1)

        self.assertEqual(len(results), 1)
        self.assertIn("support hours", results[0].chunk.text.lower())

    def test_indexer_can_query_repo_data_directory(self) -> None:
        repo_root = Path(__file__).resolve().parents[1]
        data_dir = repo_root / "data"
        with tempfile.TemporaryDirectory() as temp_dir:
            settings = Settings(
                content_data_dir=data_dir,
                retrieval_vector_store_path=Path(temp_dir) / "vectors.json",
                retrieval_top_k=2,
                retrieval_chunk_size=120,
                retrieval_chunk_overlap=20,
            )
            indexer = build_indexer(settings)
            indexer.index()

            results = run_query(indexer, "What are your support hours?", top_k=5)

        self.assertTrue(results)
        self.assertTrue(
            any("support hours" in result.chunk.text.lower() for result in results),
        )

    def test_indexer_ignores_generated_retrieval_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            self._write_json(
                data_dir / "knowledge.json",
                {"sections": {"FAQs": ["Support hours are weekdays 9 to 5."]}},
            )
            self._write_json(
                data_dir / "retrieval" / "vector_store.json",
                [{"text": "Generated vector payload should not become knowledge."}],
            )
            settings = Settings(
                content_data_dir=data_dir,
                retrieval_vector_store_path=data_dir / "retrieval" / "vector_store.json",
                retrieval_chunk_size=120,
                retrieval_chunk_overlap=20,
            )
            indexer = build_indexer(settings)

            chunks = indexer.index()

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].text, "sections > FAQs > 0: Support hours are weekdays 9 to 5.")

    def test_default_retriever_builds_missing_local_index(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            vector_store_path = data_dir / "retrieval" / "vector_store.json"
            self._write_json(
                data_dir / "business_profile.json",
                {
                    "business_name": "Glow Studio",
                    "support_hours": "Monday to Friday, 9:00 AM to 5:00 PM.",
                },
            )
            settings = Settings(
                content_data_dir=data_dir,
                retrieval_vector_store_path=vector_store_path,
                retrieval_top_k=2,
                retrieval_chunk_size=120,
                retrieval_chunk_overlap=20,
            )
            get_retriever.cache_clear()

            with patch("src.app.core.dependencies.get_config", return_value=settings):
                retriever = get_retriever()
                results = retriever.retrieve("What are your support hours?", top_k=2)

            get_retriever.cache_clear()
            index_exists = vector_store_path.exists()

            self.assertTrue(index_exists)
            self.assertTrue(results)
            self.assertTrue(
                any("support hours" in result.chunk.text.lower() for result in results),
            )

    def test_refund_query_matches_plural_refunds_policy_confidently(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            vector_store_path = data_dir / "retrieval" / "vector_store.json"
            self._write_json(
                data_dir / "knowledge.json",
                {
                    "sections": {
                        "Policies": [
                            (
                                "Escalate to a human support team when the request needs "
                                "account access, refunds, policy exceptions, or other "
                                "actions you cannot complete directly."
                            ),
                        ],
                    },
                },
            )
            settings = Settings(
                content_data_dir=data_dir,
                retrieval_vector_store_path=vector_store_path,
                retrieval_top_k=3,
                retrieval_chunk_size=200,
                retrieval_chunk_overlap=20,
            )
            indexer = build_indexer(settings)
            indexer.index()
            retriever = VectorStoreRetriever(
                embedder=indexer._embedder,
                vector_store=JsonVectorStore(vector_store_path),
                default_top_k=3,
            )
            pipeline = RetrievalPipeline(retriever=retriever)

            decision = pipeline.run("Can you help me with refund?")

        self.assertFalse(decision.used_fallback)
        self.assertEqual(decision.decision_reason, "high_confidence")
        self.assertGreaterEqual(decision.confidence_score, 0.6)
        self.assertIn("refunds", decision.retrieved_context[0].text)

    def _write_json(self, path: Path, payload: dict[str, object]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(payload), encoding="utf-8")
